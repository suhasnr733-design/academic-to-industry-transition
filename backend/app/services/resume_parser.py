# backend/app/services/resume_parser.py

import re
import PyPDF2
import docx
from typing import List, Dict, Any
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        nlp = None
except ImportError:
    spacy = None
    nlp = None
import nltk

class ResumeParser:
    """Parse resumes and extract structured information"""
    
    def __init__(self):
        # Predefined skill sets
        self.skill_keywords = {
            'Programming Languages': [
                'Python', 'Java', 'C++', 'C#', 'C', 'JavaScript', 'TypeScript',
                'Ruby', 'Swift', 'Kotlin', 'Go', 'Golang', 'Rust', 'PHP', 'R', 'Scala', 'Dart'
            ],
            'Web Technologies': [
                'HTML', 'HTML5', 'CSS', 'CSS3', 'React', 'React.js', 'Angular',
                'Vue.js', 'Vue', 'Node.js', 'Express', 'Express.js', 'Django',
                'Flask', 'FastAPI', 'Spring', 'Spring Boot', 'Hibernate', 'Next.js', 'Tailwind CSS',
                'Tailwind', 'Bootstrap', 'REST API', 'RESTful APIs', 'GraphQL', 'Redux', 'Twilio'
            ],
            'Database': [
                'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Oracle', 'Redis',
                'Firebase', 'Elasticsearch', 'SQLite', 'Cassandra', 'DynamoDB', 'DBMS', 'Pinecone', 'Vector Database'
            ],
            'Core CS Concepts': [
                'Data Structures', 'Algorithms', 'DSA', 'OOP', 'Object Oriented Programming',
                'DBMS', 'Operating Systems', 'System Design', 'Computer Networks'
            ],
            'ML/AI': [
                'Machine Learning', 'Deep Learning', 'Generative AI', 'GenAI',
                'LLM', 'LLMs', 'NLP', 'Natural Language Processing', 'Computer Vision',
                'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'HuggingFace', 'Hugging Face',
                'OpenAI', 'LangChain', 'Pinecone', 'ElevenLabs', 'RAG', 'Vector DB', 'Embeddings'
            ],
            'Cloud & DevOps': [
                'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins',
                'DevOps', 'CI/CD', 'Terraform', 'Ansible', 'Linux'
            ],
            'Data Science': [
                'Data Science', 'Data Analysis', 'Data Visualization', 'Statistics',
                'Tableau', 'Power BI', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn'
            ],
            'Tools': [
                'Git', 'GitHub', 'GitLab', 'Bitbucket', 'JIRA', 'Postman',
                'VS Code', 'IntelliJ', 'Excel', 'Figma'
            ],
            'Soft Skills': [
                'Communication', 'Leadership', 'Teamwork', 'Problem Solving',
                'Critical Thinking', 'Time Management'
            ]
        }
        
        # Flatten and sort skill list for greedy matching (longer terms first)
        all_skills_flat = [skill for skills in self.skill_keywords.values() for skill in skills]
        self.all_skills = sorted(list(set(all_skills_flat)), key=lambda s: len(s), reverse=True)
        
        # Canonical normalization map
        self.alias_map = {
            'react.js': 'React',
            'reactjs': 'React',
            'node.js': 'Node.js',
            'nodejs': 'Node.js',
            'express.js': 'Express.js',
            'expressjs': 'Express.js',
            'express': 'Express.js',
            'golang': 'Go',
            'html5': 'HTML',
            'css3': 'CSS',
            'restful apis': 'REST API',
            'rest apis': 'REST API',
            'dsa': 'Data Structures',
            'object oriented programming': 'OOP',
            'genai': 'Generative AI',
            'llms': 'LLM',
            'natural language processing': 'NLP',
            'langchain': 'LangChain',
            'pinecone': 'Pinecone',
            'openai': 'OpenAI',
            'elevenlabs': 'ElevenLabs',
            'twilio': 'Twilio',
            'hugging face': 'Hugging Face',
            'huggingface': 'Hugging Face',
            'rag': 'RAG',
            'fastapi': 'FastAPI',
            'spring boot': 'Spring Boot',
            'hibernate': 'Hibernate',
            'tailwind': 'Tailwind CSS',
            'tailwind css': 'Tailwind CSS'
        }
        
        # Education keywords
        self.education_patterns = [
            r'(B\.E|B\.Tech|M\.E|M\.Tech|B\.Sc|M\.Sc|BCA|MCA|PhD|B\.A|M\.A|MBA)\s+[^\n]*',
            r'(Bachelor|Master|Doctorate|M\.Phil)\s+(?:of|in|of Science|of Arts)\s+[^\n]*'
        ]
        
        # Experience patterns
        self.experience_patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?|Years?)\s+(?:of\s+)?(?:experience|Experience)',
            r'Experience\s*:\s*(\d+)\+?\s*(?:years?|yrs?)'
        ]

    def clean_extracted_text(self, text: str) -> str:
        """Pre-processes raw text extracted from PDF to fix OCR/glyph glitches and merged words"""
        if not text:
            return ""
        
        # 1. Clean corrupted icon font glyphs in contact header
        text = re.sub(r'\b(dphone|envel~pe|linkedinLinkedIn|githubGitHub|codeLeetCode)\b', ' ', text, flags=re.I)
        
        # 2. Normalize strange bullet symbols (e.g. 'º:' -> '•')
        text = re.sub(r'[\u00BA\u00B0\u25CF\u25CB\u25AA\u25AB]\s*:\s*', '• ', text)
        
        # 3. Separate words merged with numbers/percentages (e.g. "by70%" -> "by 70%", "90%response" -> "90% response")
        text = re.sub(r'([a-zA-Z])(\d+%)', r'\1 \2', text)
        text = re.sub(r'(\d+%)(\w)', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z])(\d{4})', r'\1 \2', text)
        
        # 4. De-merge attached lowercase-to-uppercase words (e.g. "CollegeMangaluru" -> "College Mangaluru", "SystemGitHub" -> "System GitHub")
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # 5. Clean multi-spaces while preserving line breaks
        cleaned_lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.splitlines()]
        return "\n".join(cleaned_lines)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return self.clean_extracted_text(text)

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
        return self.clean_extracted_text(text)

    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text based on file type"""
        if file_type.lower() == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path)
        elif file_type.lower() == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return self.clean_extracted_text(f.read())
            except Exception as e:
                print(f"Error reading TXT: {e}")
                return ""
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def extract_candidate_name(self, text: str) -> str:
        """Extract candidate name from the top header of the resume"""
        if not text:
            return None
        
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return None
        
        # Check first 5 lines for a plausible person name
        for line in lines[:5]:
            # Skip if contains email, phone, web, linkedin, github, or typical section titles
            if re.search(r'(@|phone|tel|\+?\d{10}|https?:|www\.|linkedin|github|leetcode|hackerrank|portfolio|curriculum|resume|summary|education|experience|skills|projects)', line, re.IGNORECASE):
                continue
            
            # Clean up punctuation and digits
            cleaned = re.sub(r'[^a-zA-Z\s\.\-]', '', line).strip()
            words = cleaned.split()
            # Supports names with initials (e.g. "Abhinandan H L")
            if 2 <= len(words) <= 5 and any(len(w) >= 3 for w in words):
                # Filter out pure headings or uppercase acronym strings like "COMPUTER SCIENCE ENGINEERING"
                if re.search(r'\b(computer|science|engineering|technology|developer|student|bachelor|master|curriculum|vitae|profile)\b', cleaned, re.IGNORECASE):
                    continue
                return cleaned.title()
        
        return None

    def extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from text using NLP and keyword matching"""
        found_skills = set()
        
        # Filter out letter-spaced uppercase header artifacts (e.g., 'S U M M A R Y')
        clean_lines = []
        for line in text.splitlines():
            tokens = line.strip().split()
            if len(tokens) >= 3 and all(len(t) == 1 for t in tokens):
                continue
            clean_lines.append(line)
        cleaned_text = "\n".join(clean_lines)

        # Method 1: Keyword matching with contextual word boundaries
        for skill in self.all_skills:
            if skill in ['C', 'R']:
                # Safely match single-letter programming languages like C or R in language lists/skill sections
                pattern = r'(?i)(?:(?:programming|languages|skills|technologies|tools)\s*:[^\n]*\b' + re.escape(skill) + r'\b)|(?:,\s*' + re.escape(skill) + r'\s*,)|(?:,\s*' + re.escape(skill) + r'\s*$)|(?:^\s*' + re.escape(skill) + r'\s*,)|(?:\b' + re.escape(skill) + r'\s*\/)|(?:\/\s*' + re.escape(skill) + r'\b)'
                if re.search(pattern, cleaned_text):
                    canonical = self.alias_map.get(skill.lower(), skill)
                    found_skills.add(canonical)
            else:
                # Whole word match for regular keywords
                pattern = r'(?i)\b' + re.escape(skill) + r'(?:\b|$)'
                if re.search(pattern, cleaned_text):
                    canonical = self.alias_map.get(skill.lower(), skill)
                    found_skills.add(canonical)
        
        # Method 2: spaCy NER for organizations and products
        if nlp is not None:
            try:
                doc = nlp(cleaned_text)
                for ent in doc.ents:
                    if ent.label_ in ["ORG", "PRODUCT"] and ent.text in self.all_skills:
                        canonical = self.alias_map.get(ent.text.lower(), ent.text)
                        found_skills.add(canonical)
            except Exception:
                pass
        
        # Method 3: Extract from bullet points using patterns
        skill_patterns = [
            r'(?:Skills|Technical Skills|Technologies|Languages|Tools|Core Concepts|Programming|Web Development|AI/ML|Database)\s*[:\-]\s*([^\n]+)',
            r'(?:Proficient|Experienced|Expertise)\s+in\s+([^\n.]+)'
        ]
        for pattern in skill_patterns:
            matches = re.finditer(pattern, cleaned_text, re.IGNORECASE)
            for match in matches:
                skills_text = match.group(1)
                for skill in self.all_skills:
                    if skill in ['C', 'R']:
                        c_pattern = r'(?i)(?:^|[\s,;:\(\[\/\-])' + re.escape(skill) + r'(?:[\s,;:\)\]\/\-]|$)'
                        if re.search(c_pattern, skills_text):
                            canonical = self.alias_map.get(skill.lower(), skill)
                            found_skills.add(canonical)
                    elif re.search(r'(?i)\b' + re.escape(skill) + r'(?:\b|$)', skills_text):
                        canonical = self.alias_map.get(skill.lower(), skill)
                        found_skills.add(canonical)
        
        return sorted(list(found_skills))

    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education details from resume with specialization, CGPA, and year range"""
        education_list = []
        
        # 1. Degree pattern with specialization (e.g. B.E. in Artificial Intelligence and Machine Learning)
        degree_patterns = [
            r'(B\.E\.?|B\.Tech\.?|Bachelor of Engineering|Bachelor of Technology|M\.E\.?|M\.Tech\.?|BCA|MCA|B\.Sc|M\.Sc|PhD)\s+(?:in\s+)?([^\n;\.,]+)',
            r'(B\.E|B\.Tech|M\.E|M\.Tech|B\.Sc|M\.Sc|BCA|MCA|PhD|B\.A|M\.A|MBA)\s+[^\n]*'
        ]
        degree_found = None
        for pattern in degree_patterns:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                if len(m.groups()) >= 2:
                    degree_found = f"{m.group(1).strip()} in {m.group(2).strip()}"
                else:
                    degree_found = m.group(0).strip()
                break

        # 2. Institution pattern (clean line bounded)
        inst_pattern = r'([A-Z][a-zA-Z\s]{2,40}(?:Engineering College|University|Institute of Technology|College of Engineering|College))'
        inst_match = re.search(inst_pattern, text)
        inst_found = inst_match.group(1).strip() if inst_match else None
        if inst_found and '\n' in inst_found:
            inst_found = inst_found.split('\n')[-1].strip()

        # 3. CGPA / GPA pattern
        gpa_pattern = r'(?:CGPA|GPA)[\s:]*(\d+\.?\d*)\s*(?:\/\s*10)?'
        gpa_match = re.search(gpa_pattern, text, re.IGNORECASE)
        gpa_found = f"CGPA: {gpa_match.group(1)} / 10" if gpa_match else None

        # 4. Years / Date range pattern
        years_pattern = r'(\b20\d{2}\s*[-–—]\s*(?:20\d{2}|Present)\b)'
        years_match = re.search(years_pattern, text, re.IGNORECASE)
        year_found = years_match.group(1) if years_match else "2023 - 2027"

        if degree_found or inst_found:
            education_list.append({
                'degree': degree_found or "Bachelor of Engineering",
                'institution': inst_found or "Engineering Institution",
                'gpa': gpa_found,
                'year': year_found
            })
        
        return education_list

    def extract_experience(self, text: str) -> Dict[str, Any]:
        """Extract work experience information including internships"""
        experience_data = {
            'years': None,
            'months': None,
            'companies': [],
            'roles': [],
            'internships': [],
            'description': []
        }
        
        # Extract years of experience
        for pattern in self.experience_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    years = int(re.search(r'\d+', str(match)).group())
                    if experience_data['years'] is None or years > experience_data['years']:
                        experience_data['years'] = years
                except:
                    pass
        
        # Extract company names (e.g. Truck Hai Technologies)
        companies = []
        roles = []
        for line in text.splitlines():
            line_str = line.strip()
            if not line_str or re.search(r'^(experience|work experience|employment|projects|education|skills)\b', line_str, re.I):
                continue
            
            # Company match
            c_match = re.search(r'\b([A-Z][a-zA-Z0-9\s]{2,35}(?:Technologies|Tech|Solutions|Labs|Systems|Services|Logistics|Private|Limited|Ltd|Inc|Corp))\b', line_str)
            if c_match:
                companies.append(c_match.group(1).strip())
            
            # Role & Intern match
            r_match = re.search(r'\b([A-Z][a-zA-Z\s]{2,30}(?:Developer|Engineer|Analyst|Scientist|Manager|Consultant|Architect)(?:\s+Intern)?)\b', line_str)
            if r_match:
                roles.append(r_match.group(1).strip())

        experience_data['companies'] = list(set([c for c in companies if len(c) > 3]))
        experience_data['roles'] = list(set([r for r in roles if len(r) > 3]))
        
        return experience_data

    def extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract project information, technologies, and quantified metrics from resume"""
        projects = []
        
        # Find project section - cleanly bounded before Technical Skills, Certifications, Achievements, Publications etc.
        project_section_patterns = [
            r'(?:PROJECTS|ACADEMIC PROJECTS|PERSONAL PROJECTS|PROJECT WORK)[\s\S]*?(?=\n\s*(?:TECHNICAL SKILLS|SKILLS|CERTIFICATIONS|COURSES|EXPERIENCE|EDUCATION|ACHIEVEMENTS|AWARDS|HONORS|PUBLICATIONS|RESEARCH|\Z))',
            r'(?:Projects|Academic Projects|Personal Projects)[\s\S]*?(?=\n\n\n|\Z)'
        ]
        
        project_text = None
        for pattern in project_section_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                project_text = match.group(0)
                break
        
        if not project_text:
            return projects

        action_verbs = r'^(?:developed|built|implemented|automated|optimized|designed|engineered|maintained|tested|integrated|created|conducted|trained|deployed|collaborated|improved|spearheaded|utilized|used)\b'

        current_project = None
        lines = [l.strip() for l in project_text.splitlines() if l.strip()]
        
        for line in lines[1:]:  # Skip header line
            clean_line = re.sub(r'^[•●○\-\*º:\s]+', '', line).strip()
            if not clean_line:
                continue

            # Check if line looks like an action verb bullet point or tech stack
            is_bullet_action = bool(re.search(action_verbs, clean_line, re.IGNORECASE))
            is_tech_header = bool(re.search(r'^(?:technologies|tech stack|tools|built with|stack)\s*[:\-]', clean_line, re.IGNORECASE))
            is_comma_tech_list = (len(clean_line.split(',')) >= 2) and any(
                skill.lower() in clean_line.lower() 
                for skill in ['python', 'flask', 'django', 'react', 'java', 'sql', 'postgresql', 'langchain', 'pinecone', 'openai', 'bootstrap', 'sqlite', 'node', 'aws', 'docker']
            )
            
            # If line is short, doesn't start with action verb, and isn't a tech list -> it's a new Project Title
            is_new_title = (
                not is_bullet_action and 
                not is_tech_header and 
                not is_comma_tech_list and
                len(clean_line) < 75 and 
                len(clean_line.split()) <= 10 and
                not clean_line.endswith('.')
            )

            if is_new_title:
                if current_project and len(current_project['title']) >= 3:
                    projects.append(current_project)
                
                clean_title = re.sub(r'\b(GitHub|Git Hub|Link|Demo|Code|Live|Repository)\b|[:\-\|\[\]\(\)]', ' ', clean_line, flags=re.IGNORECASE)
                clean_title = " ".join(clean_title.split()).strip()
                current_project = {
                    'title': clean_title,
                    'raw_block': [clean_line],
                    'description': '',
                    'technologies': [],
                    'metrics': [],
                    'duration': ''
                }
            elif current_project:
                current_project['raw_block'].append(clean_line)

        if current_project and len(current_project['title']) >= 3:
            projects.append(current_project)

        # Post-process extracted projects for technologies and metrics
        for p in projects:
            full_block = " ".join(p.get('raw_block', []))
            # Technologies
            tech_found = []
            for skill in self.all_skills:
                if re.search(r'\b' + re.escape(skill) + r'\b', full_block, re.IGNORECASE):
                    tech_found.append(self.alias_map.get(skill.lower(), skill))
            p['technologies'] = list(set(tech_found))[:6]
            
            # Metrics (% or +)
            p['metrics'] = re.findall(r'\b(?:\d+%\s*(?:reduction|increase|accuracy|efficiency|response|latency|performance)?|\d+\+\s*(?:users|queries|requests)?)\b', full_block, re.IGNORECASE)
            
            # Duration if present
            duration_match = re.search(r'(\d+)\s*(?:months?|mths?|weeks?)', full_block, re.IGNORECASE)
            p['duration'] = duration_match.group(0) if duration_match else ''

            # Description
            desc_lines = p.get('raw_block', [])[1:]
            p['description'] = " ".join(desc_lines) if desc_lines else full_block
            p.pop('raw_block', None)

        return projects

    def extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extracts candidate email, phone number, location, and bio-data clutter check"""
        email_match = re.search(r'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)', text)
        email = email_match.group(1).strip() if email_match else None
        
        phone_match = re.search(r'((?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})', text)
        phone = phone_match.group(1).strip() if phone_match else None
        
        loc_match = re.search(r'\b([A-Z][a-zA-Z\s]{2,20},\s*(?:Karnataka|Maharashtra|Tamil Nadu|Delhi|Telangana|India|USA|UK|Canada|[A-Z]{2}))\b', text)
        location = loc_match.group(1).strip() if loc_match else None
        
        has_clutter = bool(re.search(r'\b(date of birth|dob|father\'?s? name|marital status|nationality|religion|caste)\b', text, re.IGNORECASE))
        
        return {
            'candidate_name': self.extract_candidate_name(text),
            'email': email,
            'phone': phone,
            'location': location,
            'has_bio_data_clutter': has_clutter
        }

    def extract_links(self, text: str) -> Dict[str, str]:
        """Extracts multi-domain platform links"""
        patterns = {
            'github': r'(?:https?:\/\/)?(?:www\.)?github\.com\/([a-zA-Z0-9_\-\.]+)',
            'gitlab': r'(?:https?:\/\/)?(?:www\.)?gitlab\.com\/([a-zA-Z0-9_\-\.]+)',
            'linkedin': r'(?:https?:\/\/)?(?:www\.)?linkedin\.com\/(?:in|profile)\/([a-zA-Z0-9_\-\.]+)',
            'leetcode': r'(?:https?:\/\/)?(?:www\.)?leetcode\.com\/(?:u\/)?([a-zA-Z0-9_\-\.]+)',
            'hackerrank': r'(?:https?:\/\/)?(?:www\.)?hackerrank\.com\/(?:profile\/)?([a-zA-Z0-9_\-\.]+)',
            'codeforces': r'(?:https?:\/\/)?(?:www\.)?codeforces\.com\/profile\/([a-zA-Z0-9_\-\.]+)',
            'codechef': r'(?:https?:\/\/)?(?:www\.)?codechef\.com\/users\/([a-zA-Z0-9_\-\.]+)',
            'kaggle': r'(?:https?:\/\/)?(?:www\.)?kaggle\.com\/([a-zA-Z0-9_\-\.]+)',
            'huggingface': r'(?:https?:\/\/)?(?:www\.)?huggingface\.co\/([a-zA-Z0-9_\-\.]+)',
            'behance': r'(?:https?:\/\/)?(?:www\.)?behance\.net\/([a-zA-Z0-9_\-\.]+)',
            'portfolio': r'(?:https?:\/\/)?(?:[a-zA-Z0-9_\-]+\.)+(?:dev|io|me|app|vercel\.app|netlify\.app|render\.com)'
        }
        links = {}
        for platform, pattern in patterns.items():
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                links[platform] = m.group(0).strip()
        return links

    def extract_summary(self, text: str) -> str:
        """Extracts 2-3 line professional summary or career objective"""
        m = re.search(r'(?:SUMMARY|PROFESSIONAL SUMMARY|CAREER OBJECTIVE|ABOUT ME)[\s\S]*?(?=\n\s*(?:TECHNICAL SKILLS|SKILLS|EXPERIENCE|EDUCATION|PROJECTS|\Z))', text, re.IGNORECASE)
        if m:
            lines = [l.strip() for l in m.group(0).splitlines() if l.strip()]
            if len(lines) > 1:
                return " ".join(lines[1:4])
        return ""

    def extract_action_verbs_and_metrics(self, text: str) -> Dict[str, Any]:
        """Evaluates Google XYZ & STAR metric impact"""
        verbs = ['architected', 'automated', 'engineered', 'optimized', 'developed', 'spearheaded', 'deployed', 'built', 'reduced', 'increased', 'streamlined', 'integrated']
        lines = [l.strip().lower() for l in text.splitlines() if l.strip()]
        
        verb_count = sum(1 for line in lines if any(line.startswith(v) or f"• {v}" in line or f"- {v}" in line for v in verbs))
        metrics = re.findall(r'\b(?:\d+%\s*(?:reduction|increase|latency|efficiency|accuracy)?|\d+\+\s*(?:users|queries|requests)|\$[\d,]+)\b', text, re.IGNORECASE)
        
        return {
            'verb_count': verb_count,
            'metric_count': len(metrics),
            'metrics_samples': metrics[:5]
        }

    def extract_achievements(self, text: str) -> List[str]:
        """Extracts hackathons, honors, and contest rankings"""
        achievements = []
        m = re.search(r'(?:ACHIEVEMENTS|HONORS|AWARDS|HACKATHONS|COMPETITIVE PROGRAMMING)[\s\S]*?(?=\n\s*(?:PROJECTS|SKILLS|EXPERIENCE|EDUCATION|CERTIFICATIONS|PUBLICATIONS|RESEARCH|\Z))', text, re.IGNORECASE)
        if m:
            for line in m.group(0).splitlines()[1:]:
                clean = re.sub(r'^[•●○\-\*º:\s]+', '', line).strip()
                if len(clean) > 8:
                    achievements.append(clean)
        return achievements[:6]

    def extract_publications(self, text: str) -> List[str]:
        """Extracts research papers, conferences, and patents"""
        publications = []
        m = re.search(r'(?:PUBLICATIONS|RESEARCH PAPERS|PATENTS|CONFERENCES)[\s\S]*?(?=\n\s*(?:PROJECTS|SKILLS|EXPERIENCE|EDUCATION|CERTIFICATIONS|ACHIEVEMENTS|AWARDS|\Z))', text, re.IGNORECASE)
        if m:
            for line in m.group(0).splitlines()[1:]:
                clean = re.sub(r'^[•●○\-\*º:\s]+', '', line).strip()
                if len(clean) > 10:
                    publications.append(clean)
        return publications[:5]

    def extract_certifications(self, text: str) -> List[Dict[str, str]]:
        """Extract certification names and issuing organizations from resume"""
        certifications = []
        
        # Match Certifications section
        cert_match = re.search(
            r'(?:CERTIFICATIONS|CERTIFICATES|COURSES & CERTIFICATIONS|COURSES\s*:\s*)[\s\S]*', 
            text, re.IGNORECASE
        )
        if not cert_match:
            return certifications
            
        cert_text = cert_match.group(0)
        # Trim if another major section follows
        cert_text = re.split(r'\n\s*(?:PROJECTS|TECHNICAL SKILLS|EXPERIENCE|EDUCATION|SKILLS|ACHIEVEMENTS|PUBLICATIONS)\b', cert_text, flags=re.IGNORECASE)[0]
        lines = [line.strip() for line in cert_text.splitlines() if line.strip()]
        
        for line in lines[1:]:  # Skip header
            clean_line = re.sub(r'^[•●○\-\*º:\s]+', '', line).strip()
            if len(clean_line) > 4 and not re.search(r'^(certifications|certificates|courses)\b', clean_line, re.I):
                parts = re.split(r'\s*[-–—]\s*', clean_line, maxsplit=1)
                name = parts[0].strip()
                issuer = parts[1].strip() if len(parts) > 1 else 'Industry Certified'
                certifications.append({
                    'name': name,
                    'issuer': issuer,
                    'full_title': clean_line
                })
                
        return certifications

    def parse_resume(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Main method to parse resume and extract all information"""
        try:
            # Extract text
            text = self.extract_text(file_path, file_type)
            
            if not text or len(text.strip()) < 50:
                return {
                    'error': 'Could not extract text from file. File might be empty or image-based.',
                    'success': False
                }
            
            skills = self.extract_skills(text)
            category_count = sum(1 for cat, kw_list in self.skill_keywords.items() if any(k in skills for k in kw_list))
            
            # Extract all sections
            result = {
                'success': True,
                'candidate_name': self.extract_candidate_name(text),
                'personal_info': self.extract_personal_info(text),
                'links': self.extract_links(text),
                'summary': self.extract_summary(text),
                'skills': skills,
                'category_count': category_count,
                'education': self.extract_education(text),
                'experience': self.extract_experience(text),
                'projects': self.extract_projects(text),
                'certifications': self.extract_certifications(text),
                'achievements': self.extract_achievements(text),
                'publications': self.extract_publications(text),
                'metrics_analysis': self.extract_action_verbs_and_metrics(text),
                'raw_text': text,
                'text_preview': text[:500] + '...' if len(text) > 500 else text
            }
            
            # Add statistics
            result['stats'] = {
                'skill_count': len(result['skills']),
                'education_count': len(result['education']),
                'project_count': len(result['projects']),
                'certification_count': len(result['certifications']),
                'total_words': len(text.split())
            }
            
            return result
            
        except Exception as e:
            return {
                'error': str(e),
                'success': False
            }

    def parse_resume_text(self, text: str) -> Dict[str, Any]:
        """Parse raw resume text directly"""
        clean_text = self.clean_extracted_text(text)
        skills = self.extract_skills(clean_text)
        category_count = sum(1 for cat, kw_list in self.skill_keywords.items() if any(k in skills for k in kw_list))
        return {
            'success': True,
            'candidate_name': self.extract_candidate_name(clean_text),
            'personal_info': self.extract_personal_info(clean_text),
            'links': self.extract_links(clean_text),
            'summary': self.extract_summary(clean_text),
            'skills': skills,
            'category_count': category_count,
            'education': self.extract_education(clean_text),
            'experience': self.extract_experience(clean_text),
            'projects': self.extract_projects(clean_text),
            'certifications': self.extract_certifications(clean_text),
            'achievements': self.extract_achievements(clean_text),
            'publications': self.extract_publications(clean_text),
            'metrics_analysis': self.extract_action_verbs_and_metrics(clean_text),
            'raw_text': clean_text
        }